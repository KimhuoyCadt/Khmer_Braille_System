
from KhmerToBrailleLogic import khmer_to_braille
from BrailleToKhmerLogic import translateBrailleToKhmer
from FileFormat import read_ocr_file, read_text_file, read_word_file, write_doc_file, write_pdf_file, write_txt_file
# from KhmerToBrailleLogic import createUnicode
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from docx import Document
import requests, os
from PIL import Image
import fitz
import pytesseract
from khmernltk import word_tokenize
from gcc_segmentation import GCCSegmentation
from flask_cors import CORS
import logging

app = Flask(__name__)
CORS(app)
CORS(app, resources={r"/braille": {"origins": "http://127.0.0.1:5173"}})
segmenter = GCCSegmentation()
def format_khmer_and_braille_text(khmer_words, max_line_length=42):
    formatted_lines = []
    current_braille_line = ""

    for word in khmer_words:
        # Split word on \n to handle new lines within the word
        split_words = word.split('\n')
        for i, part in enumerate(split_words):
            braille_word = khmer_to_braille(part)

            # Check if adding this braille_word would exceed the max_line_length
            if len(current_braille_line) + len(braille_word) <= max_line_length:
                current_braille_line += braille_word
            else:
                formatted_lines.append(current_braille_line)
                current_braille_line = braille_word

            # If not the last part, it means there was a '\n' — force a new line
            if i < len(split_words) - 1:
                formatted_lines.append(current_braille_line)
                current_braille_line = ""

    # Add any remaining braille text to the formatted lines
    if current_braille_line:
        formatted_lines.append(current_braille_line)

    return formatted_lines

# handle for upload text file to translate
@app.route('/updateTextFile', methods=['POST'])
def updateTextFile():   
    global uploaded_text
    if request.method == 'POST':
        try:
            if 'file' not in request.files:
                return jsonify({"error": "No file part"}), 400

            upload = request.files['file']
            if upload.filename == '':
                return jsonify({"error": "No selected file"}), 400

            if upload:
                file_path = os.path.join("uploads", secure_filename(upload.filename))
                upload.save(file_path)
                print(f"File saved at: {file_path}")
            file_extension = os.path.splitext(file_path)[1]

            if(file_extension == ".txt"):
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        combined_lines = ""
                        for line in file:
                            combined_lines += line.strip() + "\n"
                            uploaded_text = combined_lines
                            print("text:", uploaded_text)
                except FileNotFoundError:
                    return 'File not found', 404
                except Exception as e:
                    return f'An error occurred: {str(e)}', 500
                
            if(file_extension ==".docx" or file_extension ==".doc"):
                if file_path:
                    try:
                        document = Document(file_path)
                        uploaded_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                        print("docx: ",uploaded_text)
                    except Exception as e:
                        error_message = f"Error reading Word file: {e}"

            if file_extension == ".pdf":
                try:
                    combined_text = ""
                    paragraph = []
                    pdf_document = fitz.open(file_path)
                    images = []
                    KHSYMBOL = set(u')(][!@#$%^&*{}+-=ˈ_.,|\/')
                    for page_number in range(pdf_document.page_count):
                        page = pdf_document.load_page(page_number)
                        pix = page.get_pixmap()
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        images.append(img)
                    pdf_document.close()
                    ocr_text = ""
                    for img in images:
                        text = pytesseract.image_to_string(img, lang="khm")  # Specify Khmer language
                        ocr_text += text
                        paragraph.append(ocr_text)
                    clean_paragraph = []
                    for para in paragraph:
                        sentences = para.split("\n")
                        for sentence in sentences:
                            if sentence == '':
                                continue
                            else:
                                clean_sentence = ''
                                for word in word_tokenize(sentence):
                                    if word in KHSYMBOL:
                                        continue
                                    else:
                                        clean_sentence += word 
                            clean_paragraph.append(clean_sentence)
                    for i in clean_paragraph:
                        combined_text += i 
                    uploaded_text = combined_text
                    print("pdf: ", uploaded_text)
                except Exception as e:
                    error_message = f"Error reading PDF file: {e}"
                    return jsonify({"error pdf": error_message})
            
        except Exception as e:
            logging.error(f"Error processing file: {e}")
            return jsonify({"error": "Internal Server Error"}), 500
        
    khmer_words = word_tokenize(uploaded_text)
    # Format into structured lines of Braille with 42-character limit
    formatted_lines = format_khmer_and_braille_text(khmer_words, max_line_length=42)
    formatted_braille_text = "\n".join(formatted_lines)

    print("Formatted Braille:\n" + formatted_braille_text)
    return jsonify({"result": formatted_braille_text, "uploaded": uploaded_text})

#handle to upload text to translate to braille
@app.route('/text', methods=['POST'])
def handleTextToBraille():
    data = request.get_json()
    uploaded_text = data.get("text", "")
    print("Uploaded text:", uploaded_text)
    khmer_words = segmenter.seg_gcc(uploaded_text)
    # Format the Khmer words into Braille text lines
    formatted_text = format_khmer_and_braille_text(khmer_words, max_line_length=42)
    
    # Join the formatted lines into a single string
    formatted_output = "\n".join(formatted_text)
    
    # Corrected print statement for the formatted output
    print("Formatted Output:\n" + formatted_output)

    # Return the result and Unicode data as JSON
    return jsonify({"result": formatted_output})

    

#handle to upload braille file to translate 
@app.route('/updatebraille', methods=['POST'])
def updateBrailleFile():   
    global uploaded_braille
    if request.method == 'POST':
        try:
            if 'file' not in request.files:
                return jsonify({"error": "No file part"}), 400

            upload = request.files['file']
            if upload.filename == '':
                return jsonify({"error": "No selected file"}), 400

            if upload:
                file_path = os.path.join("uploads", secure_filename(upload.filename))
                upload.save(file_path)
                print(f"File saved at: {file_path}")
            file_extension = os.path.splitext(file_path)[1]

            if(file_extension == ".txt"):
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        combined_lines = ""
                        for line in file:
                            combined_lines += line.strip()
                            uploaded_braille = combined_lines
                            print("text:", uploaded_braille)
                except FileNotFoundError:
                    return 'File not found', 404
                except Exception as e:
                    return f'An error occurred: {str(e)}', 500
                
            if(file_extension ==".docx" or file_extension ==".doc"):
                if file_path:
                    try:
                        document = Document(file_path)
                        uploaded_braille = "\n".join(paragraph.text for paragraph in document.paragraphs)
                        print("docx: ",uploaded_braille)
                    except Exception as e:
                        error_message = f"Error reading Word file: {e}"

            if file_extension == ".pdf":
                try:
                    with fitz.open(file_path) as pdf_document:
                        combined_text = ""
                        for page_num in range(pdf_document.page_count):
                            page = pdf_document[page_num]
                            text = page.get_text("text", )
                            combined_text += str(text)
                        
                        uploaded_braille = combined_text
                        print("pdf: ", uploaded_braille)
                except Exception as e:
                    error_message = f"Error reading PDF file: {e}"
                    return jsonify({"error pdf": error_message})
            
        except Exception as e:
            logging.error(f"Error processing file: {e}")
            return jsonify({"error": "Internal Server Error"}), 500
        
    braille = translateBrailleToKhmer(uploaded_braille)
    print(braille)
    return jsonify({"result": braille, "uploaded": uploaded_braille})

@app.route('/braille', methods=['POST'])
def handleBrailleToText():
    try:
        data = request.get_json()
        if data is None or "braille" not in data:
            return jsonify({"error": "Invalid request data"}), 400
        uploaded_braille = data.get("braille", "")
        logging.info("Received braille: %s", uploaded_braille)
        try:
            braille = translateBrailleToKhmer(uploaded_braille)
        except Exception as e:  
            return jsonify({"error": str(e)}), 500
        print("braille:", braille)
        return jsonify({"result": braille})
    except Exception as e:
        logging.exception("Error processing request")
        return jsonify({"error": "Internal Server Error"}), 500

if __name__ == '__main__':
    app.run(debug=True)
