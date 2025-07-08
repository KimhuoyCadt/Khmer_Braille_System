from khmernltk import word_tokenize
from docx import Document
import os
from PIL import Image
import pytesseract
import fitz
from docx2pdf import convert

# ========Format for writing braille to file =========
def format_output_file(words, max_line_length=42):
    
    """
    - Formate the braille structure displayed by one line contain 42 characters.
    - This function takes a list of words and formats them into lines that do not exceed the specified maximum line length.
    - It ensures that each line is as full as possible without exceeding the limit.
    """
    formatted_list = []
    current_line = ""
    for word in words:
        if len(current_line) + len(word) + 1 <= max_line_length:
            if current_line:
                current_line += " "
            current_line += word
        else:
            formatted_list.append(current_line)
            current_line = word
    if current_line:
        formatted_list.append(current_line)

    return formatted_list


# ========Format for reading file =========
def read_ocr_file(file_path):
    paragraph = []
    pdf_document = fitz.open(file_path)
    images = []
    KHSYMBOL = set(u')(][!@#$%^&*{}+-=ˈ_.,|\/')
    for page_number in range(pdf_document.page_count):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    pdf_document.close()
    ocr_text = ""
    for img in images:
        text = pytesseract.image_to_string(img, lang="khm")  # Specify Khmer language
        ocr_text += text
        paragraph.append(ocr_text)
    clean_paragraph = []
    for para in paragraph:
        sentences = para.split("\n")
        for sentence in sentences:
            if sentence == '':
                continue
            else:
                clean_sentence = ''
                for word in word_tokenize(sentence):
                    if word in KHSYMBOL:
                        continue
                    else:
                        clean_sentence += word 
            clean_paragraph.append(clean_sentence)
    return clean_paragraph

def read_word_file(file):
    all_file_word = []
    doc = Document(file)
    for line in doc.paragraphs:
        all_file_word.append(line.text)
    return '\n' .join(all_file_word)

def read_text_file(file):
    all_line_text_file = []
    with open(file, 'r', encoding="utf-8") as f:
        for i in f:
            all_line_text_file.append(i)
    return all_line_text_file


# ========Format for writing file =========
def write_doc_file(list, out_put_file):
    doc = Document()
    for line in list:
        doc.add_paragraph(line)
        doc.save(out_put_file)
    
def write_pdf_file(list, out_put_file):
    write_doc_file(list, "output.docx")
    convert("output.docx", out_put_file)

def write_txt_file(list,out_put_file):
    with open(out_put_file, 'a', encoding='utf-8') as f:
        for line in list:
            f.write(line)
